# -*- coding: utf-8 -*-
"""生成测试题库 Word 文档（覆盖各种格式），供 update_bank 解析准确性测试使用"""
import docx


def make_test_docx(path):
    """标准测试文档：5 单选（含代码/制表符/无前缀/笔误修正）+ 5 判断（含共享答案串/末题代码）+ 简答/多选(应被跳过)"""
    doc = docx.Document()
    # 一、单选题
    doc.add_paragraph("一、单选题")
    doc.add_paragraph("1. 以下哪个不是Python关键字？（  ）")
    for o in ("A. def", "B. class", "C. import", "D. main"):
        doc.add_paragraph(o)
    doc.add_paragraph("答案：D")

    doc.add_paragraph("2. 阅读以下代码，输出结果是（  ）")
    doc.add_paragraph("int a = 5;")
    doc.add_paragraph('printf("%d", a++);')
    for o in ("A. 5", "B. 6", "C. 7", "D. 编译错误"):
        doc.add_paragraph(o)
    doc.add_paragraph("答案：A")

    doc.add_paragraph("3. 下列哪个是Linux发行版？")
    doc.add_paragraph("A. Ubuntu\tB. Windows\tC. macOS\tD. Android")
    doc.add_paragraph("答案：A")

    doc.add_paragraph("4. 以下哪个是关系型数据库？")
    for o in ("MySQL", "Oracle", "PostgreSQL", "MongoDB"):
        doc.add_paragraph(o)
    doc.add_paragraph("答案：D")

    doc.add_paragraph("5. 在Linux中，查看服务器网口配置的命令是什么？")
    for o in ("A. ipconfig", "B. show", "C. interface", "D. ifconfig"):
        doc.add_paragraph(o)
    doc.add_paragraph("答案：D")

    # 二、判断题
    doc.add_paragraph("二、判断题")
    doc.add_paragraph("1. C语言中数组下标从0开始。（  ）")
    doc.add_paragraph("2. Python是编译型语言。（  ）")
    doc.add_paragraph("3. 循环可以相互嵌套。（  ）")
    doc.add_paragraph("答案：√××")
    doc.add_paragraph("4. Linux命令ls可以查看文件。（  ）")
    doc.add_paragraph("5. 阅读以下代码，输出结果是1。（  ）")
    doc.add_paragraph("int x = 1;")
    doc.add_paragraph('printf("%d", x);')
    doc.add_paragraph("答案：√√")

    # 三、简答题（应被跳过）
    doc.add_paragraph("三、简答题")
    doc.add_paragraph("1. 简述进程与线程的区别。")
    doc.add_paragraph("答案：略")

    # 四、多项选择题（应被跳过，只识别单选）
    doc.add_paragraph("四、多项选择题")
    doc.add_paragraph("1. 以下哪些是编程语言？（  ）")
    for o in ("A. C", "B. Python", "C. Java", "D. 以上都是"):
        doc.add_paragraph(o)
    doc.add_paragraph("答案：ABCD")

    doc.save(path)


def make_test_docx2(path):
    """第二个题库文档：内容完全不同（1 单选 + 1 判断），用于多题库选择测试"""
    doc = docx.Document()
    doc.add_paragraph("一、单选题")
    doc.add_paragraph("1. 世界上最高的山峰是？（  ）")
    for o in ("A. 泰山", "B. 珠穆朗玛峰", "C. 华山", "D. 黄山"):
        doc.add_paragraph(o)
    doc.add_paragraph("答案：B")
    doc.add_paragraph("二、判断题")
    doc.add_paragraph("1. 地球是太阳系的行星。（  ）")
    doc.add_paragraph("答案：√")
    doc.save(path)
